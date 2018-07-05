from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import pymysql
from docx.enum.style import WD_STYLE_TYPE


# def get_db():
#     # 打开数据库连接
#     db = pymysql.connect(
#         host="192.168.121.159", user="dept_jz",
#         password="nimo_PL<", db="uat_exue_resource", port=42578,
#         charset="utf8"
#     )
#     return db

def get_db():
    # 打开数据库连接
    db = pymysql.connect(
        host="123.206.227.74", user="root",
        password="exue2017", db="sit_exue_attendance", port=3306,
        charset="utf8"
    )
    return db


db = get_db()
cur = db.cursor(pymysql.cursors.DictCursor)

document = Document()
keys = ['序号', '列名', '数据类型', '主键', '外键', '允许为空', '默认值', '说明']
null_dict = {"YES": "是", "NO": "否"}
subject = ["yw", "sx", "yy", "dl", "hx", "ls", "wl", "zz", "sw", 'kx', "sp", "dd", "ty", "ms", "mu", "ps", 'xx']


def main():
    no = 1
    cur.execute(
        "select table_name,table_comment from information_schema.tables a WHERE a.TABLE_SCHEMA = 'sit_exue_attendance';")
    for item in cur.fetchall():
        # flag = False
        # for s in subject:
        #     if s in item.get('table_name'):
        #         flag = True
        # if flag and 'yw' not in item.get('table_name') and item.get('table_name') not in ['t_official_activity',
        #                                                                                   't_res_question_type',
        #                                                                                   't_special_record']:
        #     continue
        if 't_card_checkinfo' in item.get('table_name') and '397696053941383168' not in item.get('table_name'):
            continue

        paragraph = document.add_paragraph(str(no) + ' ' + item.get('table_comment') + ' ' + item.get('table_name'))
        run = paragraph.add_run()
        run.font.size = Pt(24)
        num = cur.execute("select * from information_schema.columns where table_schema = 'sit_exue_attendance' "
                          "and table_name = '{table_name}';".format(table_name=item.get('table_name')))
        table = document.add_table(rows=num + 1, cols=8)
        hdr_cells = table.rows[0].cells
        for i in range(8):
            hdr_cells[i].text = keys[i]
        for index, data in enumerate(cur.fetchall()):
            print(data)
            hdr_cells = table.rows[index + 1].cells
            hdr_cells[0].text = str(index + 1)
            hdr_cells[1].text = data.get('COLUMN_NAME')
            hdr_cells[2].text = data.get('COLUMN_TYPE')
            if data.get('COLUMN_KEY') == 'PRI':
                hdr_cells[3].text = '是'
            else:
                hdr_cells[3].text = '否'
            hdr_cells[4].text = '否'
            hdr_cells[5].text = null_dict.get(data.get('IS_NULLABLE'))
            hdr_cells[6].text = data.get('COLUMN_DEFAULT') if data.get('COLUMN_DEFAULT') else ''
            hdr_cells[7].text = data.get('COLUMN_COMMENT') if data.get('COLUMN_COMMENT') else ''

        no += 1
        document.add_paragraph("\n")
    # document.add_page_break() # 换行
    document.save('F:/table_attendance.docx')


if __name__ == '__main__':
    main()
